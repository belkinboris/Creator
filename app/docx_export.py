"""Экспорт бизнес-плана в .docx.

G3 (PRODUCT_ROADMAP, разбор соцплан.рф владельцем): комиссии соцзащиты
сдают документ, а не ссылку на веб-страницу, а конкурент прямо продаёт
«Скачайте pdf/docx(word)». PDF у нас уже был (кнопка «Скачать PDF» на
`report.html` — печать браузером по `@media print`), файла .docx не было
вовсе. python-docx, без внешних сервисов — тот же принцип единого
FastAPI-процесса, что и весь проект (см. CLAUDE.md «Архитектурные детали»).
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt

#: Курсивом и мельче — та же формулировка, что уже стоит в FAQ на витрине
#: соцконтракта («Нет. Вы получаете обоснованный текст и цифры для
#: бизнес-плана — перенести их в бланк... нужно самостоятельно»). Документ
#: не должен обещать больше, чем обещает сайт.
_FOOTER = ("Черновик подготовлен сервисом Создатель (projectsozdatel.ru). "
          "Не является официальным документом по форме вашего региона — "
          "перенесите текст и цифры в нужный бланк самостоятельно.")


def _add_body_paragraphs(doc: Document, body: str) -> None:
    for para in str(body or "").split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)


def _add_table(doc: Document, table: dict) -> None:
    """Два вида таблиц -- смета деньгами (money) и план запуска этапами
    (stages), см. table_kind в report_engine.SECTION_SPECS (G6,
    PRODUCT_ROADMAP). Разные столбцы, у stages нет строки «Итого»."""
    rows = table.get("rows") or []
    if not rows:
        return
    caption = table.get("caption")
    if caption:
        p = doc.add_paragraph()
        p.add_run(str(caption)).bold = True
    if table.get("kind") == "stages":
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Этап", "Что сделать", "Срок", "Ответственный")
        for row in rows:
            cells = t.add_row().cells
            cells[0].text = str(row.get("stage", ""))
            cells[1].text = str(row.get("what", ""))
            cells[2].text = str(row.get("deadline", "") or "—")
            cells[3].text = str(row.get("who", "") or "—")
        return
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text = "Статья", "Сумма, ₽"
    for row in rows:
        cells = t.add_row().cells
        cells[0].text = str(row.get("item", ""))
        cells[1].text = f'{int(row.get("sum", 0)):,}'.replace(",", " ")
    total = table.get("total")
    if isinstance(total, (int, float)):
        cells = t.add_row().cells
        cells[0].text = "Итого"
        cells[1].text = f'{int(total):,}'.replace(",", " ")


def build_docx(*, doc_title: str, idea: str, core: dict, sections: list[dict]) -> bytes:
    """Собрать файл .docx из уже сгенерированного отчёта (ядро + разделы).

    Не вызывает модель и не трогает БД — только форматирует то, что уже
    посчитано и сохранено, тем же контентом, что виден на `/report/{id}`.
    """
    doc = Document()
    doc.add_heading(doc_title, level=0)
    doc.add_paragraph(idea).runs[0].italic = True

    score = core.get("viability_score")
    if isinstance(score, (int, float)):
        doc.add_heading("Оценка жизнеспособности", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"{int(score)}/100")
        run.bold = True
        run.font.size = Pt(20)
        label = core.get("viability_label")
        if label:
            doc.add_paragraph(str(label)).runs[0].bold = True
        summary = core.get("viability_summary")
        if summary:
            doc.add_paragraph(str(summary))
        risks = core.get("top_risks") or []
        if risks:
            doc.add_heading("Ключевые риски", level=2)
            for r in risks:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(r.get("title", ""))).bold = True
                body = r.get("body")
                if body:
                    doc.add_paragraph(str(body))

    for sec in sections:
        doc.add_heading(str(sec.get("title", "")), level=1)
        _add_body_paragraphs(doc, sec.get("body", ""))
        table = sec.get("table")
        if table:
            _add_table(doc, table)

    footer = doc.add_paragraph()
    run = footer.add_run(_FOOTER)
    run.italic = True
    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
